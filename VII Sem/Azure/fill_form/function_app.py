import os
import re
import tempfile
import logging
from typing import Dict, Set

import requests
import azure.functions as func
from docxtpl import DocxTemplate
from docx import Document

TEMPLATE_FILENAME = "Canara_Bank_Template.docx"

TRANSLATOR_KEY = os.getenv("TRANSLATOR_KEY")
TRANSLATOR_REGION = os.getenv("TRANSLATOR_REGION")
TRANSLATOR_ENDPOINT = os.getenv("TRANSLATOR_ENDPOINT") 

TRANSLATOR_CA_BUNDLE = os.getenv("TRANSLATOR_CA_BUNDLE")
TRANSLATOR_SKIP_SSL_VERIFY = os.getenv("TRANSLATOR_SKIP_SSL_VERIFY", "false").lower() in ("1", "true", "yes")

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")

def _build_translator_url() -> str | None:
    
    if not TRANSLATOR_ENDPOINT:
        return None
    ep = TRANSLATOR_ENDPOINT.strip()
    if "/translate" in ep:
        return ep
    return ep.rstrip("/") + "/translate"

def translate_text(kannada_text: str) -> str:

    if not kannada_text:
        return ""

    translate_url = _build_translator_url()
    if not translate_url or not TRANSLATOR_KEY:
        logging.info("Translator not configured; returning original text.")
        return kannada_text

    headers = {
        "Ocp-Apim-Subscription-Key": TRANSLATOR_KEY,
        "Content-Type": "application/json"
    }
    if TRANSLATOR_REGION:
        headers["Ocp-Apim-Subscription-Region"] = TRANSLATOR_REGION

    params = {"api-version": "3.0", "to": "en"}
    body = [{"text": kannada_text}]

    verify_val = True
    if TRANSLATOR_SKIP_SSL_VERIFY:
        verify_val = False
    elif TRANSLATOR_CA_BUNDLE:
        verify_val = TRANSLATOR_CA_BUNDLE 

    try:
        resp = requests.post(
            translate_url,
            headers=headers,
            params=params,
            json=body,
            timeout=10,
            verify=verify_val
        )
        resp.raise_for_status()
        data = resp.json()
        if data and isinstance(data, list) and "translations" in data[0]:
            return data[0]["translations"][0].get("text", kannada_text)
        return kannada_text
    except requests.exceptions.SSLError:
        logging.exception("SSL error calling translator. Check TRANSLATOR_CA_BUNDLE or TRANSLATOR_SKIP_SSL_VERIFY.")
        return kannada_text
    except Exception:
        logging.exception("Translation failed; returning original text.")
        return kannada_text

def extract_placeholders(template_path: str) -> Set[str]:

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at: {template_path}")

    try:
        doc = Document(template_path)
    except Exception as e:
        logging.exception("Failed to open template with python-docx.")
        raise RuntimeError(f"Failed to open template as a .docx. Underlying error: {e}")

    placeholders = set()

    # paragraphs
    for para in doc.paragraphs:
        for match in PLACEHOLDER_RE.finditer(para.text or ""):
            placeholders.add(match.group(1))

    # tables/cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for match in PLACEHOLDER_RE.finditer(cell.text or ""):
                    placeholders.add(match.group(1))

    return placeholders

def build_context(placeholders: Set[str], input_fields: Dict[str, str]) -> Dict[str, str]:
    context: Dict[str, str] = {}
    for ph in placeholders:
        raw_value = input_fields.get(ph, "")
        if raw_value:
            translated = translate_text(raw_value)
            context[ph] = translated
        else:
            context[ph] = ""
    return context

def render_doc(template_path: str, context: Dict[str, str]) -> bytes:
    try:
        doc = DocxTemplate(template_path)
    except Exception as e:
        logging.exception("DocxTemplate failed to open template.")
        raise RuntimeError(f"Failed to open template with DocxTemplate: {e}")

    try:
        doc.render(context)
    except Exception as e:
        logging.exception("DocxTemplate render failed.")
        raise RuntimeError(f"Failed to render template: {e}")

    # Save to temp file and return bytes
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        temp_path = tmp.name

    try:
        doc.save(temp_path)
        with open(temp_path, "rb") as f:
            return f.read()
    finally:
        try:
            os.remove(temp_path)
        except Exception:
            pass

def fill_word_main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("fill_word_main invoked")

    try:
        body = req.get_json()
    except Exception:
        return func.HttpResponse("Invalid JSON", status_code=400)

    if not body or "fields" not in body:
        return func.HttpResponse("JSON must contain 'fields' object", status_code=400)

    fields_input = body["fields"]
    if not isinstance(fields_input, dict):
        return func.HttpResponse("'fields' must be an object/dictionary", status_code=400)

    base_dir = os.path.dirname(__file__)
    template_path = os.path.join(base_dir, TEMPLATE_FILENAME)
    logging.info(f"Looking for template at: {template_path} (exists={os.path.exists(template_path)})")

    if not os.path.exists(template_path):
        return func.HttpResponse(f"Template not found: {template_path}", status_code=500)

    try:
        placeholders = extract_placeholders(template_path)
        logging.info(f"Placeholders found: {placeholders}")
        context = build_context(placeholders, fields_input)
        logging.info(f"Rendering template with context: { {k: (v[:40] + '...' if len(v)>40 else v) for k,v in context.items()} }")
        file_bytes = render_doc(template_path, context)
    except Exception as e:
        logging.exception("Processing error")
        return func.HttpResponse(f"Processing error: {e}", status_code=500)

    headers = {"Content-Disposition": 'attachment; filename="filled_canara_bank.docx"'}
    return func.HttpResponse(
        body=file_bytes,
        status_code=200,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )
 